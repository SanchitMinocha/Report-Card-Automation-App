import sys
import re
import pandas as pd
import docx
from tkinter.filedialog import askopenfilename,askdirectory
from tkinter import messagebox 
import tkinter as tk
import ntpath
from tkinter import *
from tkinter import ttk
import traceback

###### For error display start

class TopErrorWindow(tk.Tk):
    def __init__(self, title, message, detail):
        super().__init__()
        self.details_expanded = False
        self.title(title)
        self.geometry('400x110+500+350')
        self.minsize(350, 75)
        self.maxsize(600, 400)
        self.resizable(False, False)
        self.rowconfigure(0, weight=0)
        self.rowconfigure(1, weight=1)
        self.columnconfigure(0, weight=1)

        button_frame = tk.Frame(self)
        button_frame.grid(row=0, column=0, sticky='nsew')
        button_frame.columnconfigure(0, weight=1)
        button_frame.columnconfigure(1, weight=1)

        text_frame = tk.Frame(self)
        text_frame.grid(row=1, column=0, padx=(7, 7), pady=(7, 7), sticky='nsew')
        text_frame.rowconfigure(0, weight=1)
        text_frame.columnconfigure(0, weight=1)

        ttk.Label(button_frame, text=message,wraplength=390,font=("Arial", 10)).grid(row=0, column=0, columnspan=3, pady=(7, 7), 
                                                                  padx=(7, 7), sticky='w')
        ttk.Button(button_frame, text='OK', command=self.destroy).grid(row=1, column=1, sticky='e',pady=(25, 7))
        ttk.Button(button_frame, text='Details',
                   command=self.toggle_details).grid(row=1, column=2, padx=(7, 7), sticky='e',pady=(25, 7))

        self.textbox = tk.Text(text_frame, height=6)
        self.textbox.insert('1.0', detail)
        self.textbox.config(state='disabled')
        self.scrollb = tk.Scrollbar(text_frame, command=self.textbox.yview)
        self.textbox.config(yscrollcommand=self.scrollb.set)
        self.mainloop()

    def toggle_details(self):
        if self.details_expanded:
            self.textbox.grid_forget()
            self.scrollb.grid_forget()
            self.resizable(False, False)
            self.geometry('400x110')
            self.details_expanded = False
        else:
            self.textbox.grid(row=0, column=0, sticky='nsew')
            self.scrollb.grid(row=0, column=1, sticky='nsew')
            self.resizable(True, True)
            self.geometry('500x250')
            self.details_expanded = True

######### For error display end

excel_filetype=[("Excel Files","*.xlsx")]
word_filetype=[("Word Files","*.docx")]

def browsefile(file_textbox,filetype):
    filename = askopenfilename(filetypes=filetype)
    file_textbox.delete(0,'end')
    file_textbox.insert(0, filename)
    
def browsefolder(folder_textbox):
    folderpath = askdirectory()
    folder_textbox.delete(0,'end')
    folder_textbox.insert(0, folderpath)
    
def only_numbers(char):
    return char.isdigit()

def getInput_ReportCardGenerate():
    global result_sheet_file
    global report_card_template
    global report_card_save_folder
    global no_of_students
    global no_working_days
    result_sheet_file = result_sheet_file_path.get()
    report_card_template = card_template_file_path.get()
    report_card_save_folder = report_card_folder_path.get()
    no_of_students = no_of_students_box.get()
    no_working_days = no_working_days_box.get()
    
def generate_ReportCards():
    try:
        no_of_student=int(no_of_students)
        no_working_day=int(no_working_days)
        result=pd.read_excel(result_sheet_file,sheet_name='COMPILED RESULT')

        for student_no in range(1,no_of_student+1):
            document = docx.Document(docx = report_card_template)
            stud_data= result[result['Unnamed: 0']==student_no]

            document.tables[0].cell(3,0).paragraphs[0].add_run(' '+str(stud_data.iloc[0,1])).bold=True
            document.tables[0].cell(3,3).paragraphs[0].add_run(' '+str(student_no)).bold=True

            marks_column_counter=2
            for row_no,row in enumerate(document.tables[1].rows):
                if(row_no!=0):
                    for cell_no,cell in enumerate(row.cells):
                        if(cell_no!=0):
                            cell.text=str(stud_data.iloc[0,marks_column_counter])
                            marks_column_counter+=1

            for row_no,row in enumerate(document.tables[2].rows):
                if(row_no!=0):
                    for cell_no,cell in enumerate(row.cells):
                        if(cell_no==0):
                            cell.text=str(stud_data.iloc[0,marks_column_counter])+' / 500'
                        else:
                            cell.text=str(stud_data.iloc[0,marks_column_counter])
                        marks_column_counter+=1

            document.tables[3].cell(0,0).paragraphs[0].add_run(' '+str(no_working_day))
            document.tables[3].cell(0,1).paragraphs[0].add_run(' '+str(stud_data.iloc[0,marks_column_counter]))
            marks_column_counter+=1

            for paragraph in document.paragraphs:
                if('GENERAL REMARKS' in paragraph.text):
                    paragraph.add_run(str(stud_data.iloc[0,marks_column_counter])).underline=True

            document.save(report_card_save_folder+'/roll_no_'+str(student_no)+'.docx')
        
        messagebox.showinfo("Successful","The Report Cards have been generated successfully.")
        
    except Exception as error:
        title = 'Failure'
        message = "An error has occurred: '{}'.".format(error)
        detail = traceback.format_exc(chain=False)
        TopErrorWindow(title, message, detail)

#Opening root window
root = Tk()

##### Root window Properties
root.option_add("*Button.Background", "White")
root.option_add("*Button.Foreground", "Black")
root.title('Report Card Generation')
root.geometry("675x280") #App Window Size
root.configure(bg='#d7cdef')  
# root['background']='#856ff8'#App background color
root.minsize(580,250) #App min size
# root.resizable(0, 0) #Don't allow resizing 
text_fg_main="#255fdf"

#Adding padding between tabs
style = ttk.Style()                     
current_theme =style.theme_use()
style.theme_settings(current_theme, settings={"TNotebook.Tab": {"configure": {"padding": [5, 1]}, 
                                                               "map": { "foreground": [("selected", "#312c78"),
                                                                                       ("active", "#000000")]
                                                                      }
                                                              }
                                            })

#Creating tab menu and tab layout

tabControl = ttk.Notebook(root)
tabControl.pack(expand=0, fill=X)

report_card_tab = ttk.Frame(tabControl)

tabControl.add(report_card_tab, text='Generate Report Cards')

tk.Grid.rowconfigure(root, 0, weight=1)
tk.Grid.columnconfigure(root, 0, weight=1)
tabControl.grid(column=0, row=0, sticky=tk.E+tk.W+tk.N+tk.S)


############## Tab 1 - Project Review Report

#Adding frame 1

main2 = tk.Frame(master=report_card_tab)
main2.pack(fill=tk.BOTH, expand=1, padx=5,pady=10)

main2.columnconfigure(0, weight=2)
main2.columnconfigure(1, weight=7)
main2.columnconfigure(3, weight=1)
# back.rowconfigure(0, weight=1)
# back.rowconfigure(1, weight=1)

#1st Row
result_sheet_label = Label(master=main2, text='Result Sheet',font=("Arial", 11),fg=text_fg_main)
result_sheet_label.grid(row=0,column=0,sticky=W, padx=10,pady=3)

result_sheet_file_path=Entry(master=main2)
result_sheet_file_path.grid(row=0,column=1,columnspan=2,sticky=W+E,pady=3)

result_sheet_browsefilebutton = Button(master=main2, text="Browse", command=lambda:[browsefile(result_sheet_file_path,excel_filetype)],font=("Arial", 10))
result_sheet_browsefilebutton.grid(row=0,column=3,sticky=E,padx=10,pady=3)

#2nd Row
card_template_label = Label(master=main2, text='Report Card Template',font=("Arial", 11),fg=text_fg_main)
card_template_label.grid(row=1,column=0,sticky=W, padx=10,pady=3)

card_template_file_path=Entry(master=main2)
card_template_file_path.grid(row=1,column=1,columnspan=2,sticky=W+E,pady=3)

card_template_browsefilebutton = Button(master=main2, text="Browse", command=lambda:[browsefile(card_template_file_path,word_filetype)],font=("Arial", 10))
card_template_browsefilebutton.grid(row=1,column=3,sticky=E,padx=10,pady=3)

#3rd Row

report_card_folder_selection = Label(master=main2, text='Folder to save report card(s)',font=("Arial", 11),fg=text_fg_main)
report_card_folder_selection.grid(row=2,column=0,sticky=W, padx=10,pady=3)

report_card_folder_path=Entry(master=main2)
report_card_folder_path.grid(row=2,column=1,columnspan=2,sticky=W+E,pady=3)

report_card_browsefolderbutton = Button(master=main2, text="Browse", command=lambda:[browsefolder(report_card_folder_path)],font=("Arial", 10))
report_card_browsefolderbutton.grid(row=2,column=3,sticky=E, padx=10,pady=3)

### Defining Validation ###
validation = main2.register(only_numbers)

#4th Row
no_of_students_label = Label(master=main2, text='Total Number of Students',font=("Arial", 11),fg=text_fg_main)
no_of_students_label.grid(row=3,column=0,sticky=W, padx=10,pady=3)

no_of_students_box=Entry(master=main2, validate="key", validatecommand=(validation, '%S'))
no_of_students_box.grid(row=3,column=1,columnspan=2,sticky=W+E,pady=3)

#5th Row
no_working_days_label = Label(master=main2, text='Number of working days',font=("Arial", 11),fg=text_fg_main)
no_working_days_label.grid(row=4,column=0,sticky=W, padx=10,pady=3)

no_working_days_box=Entry(master=main2,  validate="key", validatecommand=(validation, '%S'))
no_working_days_box.grid(row=4,column=1,columnspan=2,sticky=W+E,pady=3)

#Frame 2
panel2 = tk.Frame(master=report_card_tab)
panel2.pack(fill=BOTH, expand=0, padx=5,pady=15)

panel2.columnconfigure(0, weight=50)
panel2.columnconfigure(1, weight=1)


generate_report_cards_button = Button(master=panel2, text='Generate Report Card(s)', 
                                             command=lambda:[getInput_ReportCardGenerate(),generate_ReportCards()], font=("Arial", 10))
generate_report_cards_button.grid(row=0,column=0,sticky=E, padx=10)

close = Button(master=panel2, text='Close', command=root.destroy,font=("Arial", 10))
close.grid(row=0,column=1,sticky=E, padx=10)

mainloop()
